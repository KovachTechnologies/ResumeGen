"""
Script to generate a cover letter Word document from a template file, using data from a JSON file and command-line arguments.
"""

import argparse
import datetime
import json
import logging
from pathlib import Path
from typing import Dict, Any

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.style import WD_STYLE_TYPE

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

def load_template(template_path: Path) -> str:
    """Load the cover letter template from a file."""
    try:
        with template_path.open("r") as f:
            return f.read()
    except IOError as e:
        logger.error(f"Failed to read template from {template_path}: {str(e)}")
        raise

def load_json(json_path: Path) -> Dict[str, Any]:
    """Load JSON data from a file."""
    try:
        with json_path.open("r") as f:
            return json.load(f)
    except (IOError, json.JSONDecodeError) as e:
        logger.error(f"Failed to load JSON from {json_path}: {str(e)}")
        raise

def create_cover_letter(data: Dict[str, str], template: str, output_file: str) -> None:
    """Generate a Word document from the cover letter template with provided data."""
    try:
        # Initialize Word document
        doc = Document()

        # Apply document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Set page margins
        for section in doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

        # Replace placeholders in template
        letter_content = template
        for key, value in data.items():
            placeholder = f"${key}"
            letter_content = letter_content.replace(placeholder, value)

        # Split content into lines and process
        lines = letter_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Handle specific sections with custom formatting
            if line == 'Dear Hiring Manager,':
                para = doc.add_paragraph(line)
                para.style = 'Normal'
            elif line == 'Sincerely,':
                para = doc.add_paragraph(line)
                para.style = 'Normal'
                para.space_before = Pt(12)
            elif line == data['name'] and lines[lines.index(line) - 1].strip() == 'Sincerely,':
                para = doc.add_paragraph(line)
                para.style = 'Normal'
            else:
                # Assume body paragraphs or header lines
                para = doc.add_paragraph(line)
                para.style = 'Normal'
                if line in [data['name'], data['address'], f"{data['phone']} | {data['email']}", data['date']]:
                    para.space_after = Pt(6)
                else:
                    para.space_before = Pt(12)
                    para.space_after = Pt(12)

        # Save the document
        doc.save(output_file)
        logger.info(f"Cover letter saved to {output_file}")
    except Exception as e:
        logger.error(f"Failed to generate cover letter: {str(e)}")
        raise

def main() -> None:
    """Main function to run the cover letter generator."""
    parser = argparse.ArgumentParser(
        description="Generate a cover letter Word document from a template and JSON data.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "--json",
        type=Path,
        required=True,
        help="Path to JSON file containing name, phone, email, and address in the header",
    )
    parser.add_argument(
        "--template",
        type=Path,
        default=Path("templates/cover_letter.txt"),
        help="Path to cover letter template file",
    )
    parser.add_argument(
        "--position",
        type=str,
        default="Principal Software Engineer",
        help="Position title for the cover letter",
    )
    parser.add_argument(
        "--company",
        type=str,
        required=True,
        help="Company name for the cover letter",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="{datetime}_cover_letter.docx",
        help="Output Word document path (supports {datetime} placeholder)",
    )
    args = parser.parse_args()

    # Replace datetime placeholder in output path
    output_file = args.output.format(datetime=datetime.datetime.now().strftime("%Y-%m-%d"))

    try:
        # Load template
        template = load_template(args.template)

        # Load JSON data
        json_data = load_json(args.json)
        header = json_data.get("header", {})
        
        # Validate required JSON fields
        required_fields = ["name", "phone", "email", "address"]
        for field in required_fields:
            if field not in header or not header[field]:
                raise ValueError(f"Missing or empty field '{field}' in JSON header")

        # Prepare data dictionary for template replacement
        data = {
            "name": header["name"],
            "address": header["address"],
            "phone": header["phone"],
            "email": header["email"],
            "date": datetime.datetime.now().strftime("%B %d, %Y"),  # Today's date: May 24, 2025
            "position": args.position,
            "company": args.company
        }

        # Generate cover letter
        create_cover_letter(data, template, output_file)
    except Exception as e:
        logger.critical(f"Application failed: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()
