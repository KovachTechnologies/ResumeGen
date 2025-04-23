import unittest
from unittest.mock import patch, Mock, MagicMock
from pathlib import Path
import json
import logging
import docx
import requests
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from resumegen import ResumeGen, load_json, fetch_json, logger

class TestResumeGen(unittest.TestCase):
    def setUp(self):
        """Set up test fixtures."""
        self.config = {
            "font": {"name": "Arial", "size": 8},
            "margins": {"top": 0.5, "bottom": 0.5, "left": 1.0, "right": 1.0},
        }
        self.resume_data = {
            "header": {
                "name": "Nicholas Borbaki",
                "title": "Senior Software Engineer"
            },
            "contents": [
                {
                    "title": "Experience",
                    "id": 5,
                    "content": [
                        {
                            "id": 1,
                            "position": "Senior Software Engineer, <a href='https://www.example.com'>Example</a>",
                            "date": "2021-Present",
                            "items": [
                                "Test item with <a href='https://test.com'>Test Link</a>."
                            ]
                        }
                    ]
                }
            ]
        }
        self.test_file = Path("test_resume.json")

    def test_load_json_success(self):
        """Test loading valid JSON from a file."""
        with patch("pathlib.Path.open", new_callable=MagicMock) as mock_open:
            mock_open.return_value.__enter__.return_value.read.return_value = json.dumps(self.resume_data)
            result = load_json(self.test_file)
            self.assertEqual(result, self.resume_data)
            mock_open.assert_called_once_with("r")

    def test_load_json_file_not_found(self):
        """Test loading JSON when file is not found."""
        with patch("pathlib.Path.open", side_effect=IOError("File not found")), \
             self.assertLogs("resumegen", level="ERROR") as log:
            with self.assertRaises(IOError):
                load_json(self.test_file)
            self.assertIn("Failed to load JSON from test_resume.json: File not found", log.output[0])

    def test_load_json_invalid_json(self):
        """Test loading invalid JSON."""
        with patch("pathlib.Path.open", new_callable=MagicMock) as mock_open, \
             self.assertLogs("resumegen", level="ERROR") as log:
            mock_open.return_value.__enter__.return_value.read.return_value = "invalid json"
            with self.assertRaises(json.JSONDecodeError):
                load_json(self.test_file)
            self.assertIn("Failed to load JSON from test_resume.json", log.output[0])

    def test_fetch_json_success(self):
        """Test fetching valid JSON from a URL."""
        with patch("requests.get") as mock_get:
            mock_response = Mock()
            mock_response.json.return_value = self.resume_data
            mock_response.raise_for_status.return_value = None
            mock_get.return_value = mock_response
            result = fetch_json("https://example.com/resume.json", {"User-Agent": "test"})
            self.assertEqual(result, self.resume_data)
            mock_get.assert_called_once_with("https://example.com/resume.json", headers={"User-Agent": "test"})

    def test_fetch_json_request_failure(self):
        """Test fetching JSON when the request fails."""
        with patch("requests.get") as mock_get, \
             self.assertLogs("resumegen", level="ERROR") as log:
            mock_get.side_effect = requests.RequestException("Network error")
            with self.assertRaises(requests.RequestException):
                fetch_json("https://example.com/resume.json", {"User-Agent": "test"})
            self.assertIn("Failed to fetch JSON from https://example.com/resume.json", log.output[0])

    def test_remove_hyperlink(self):
        """Test removing HTML <a> tags from text."""
        generator = ResumeGen(self.config)
        input_text = "Senior Software Engineer, <a href='https://www.example.com'>Example</a>"
        expected = "Senior Software Engineer, Example"
        result = generator.remove_hyperlink(input_text)
        self.assertEqual(result, expected)

    def test_remove_hyperlink_no_tags(self):
        """Test remove_hyperlink with no <a> tags."""
        generator = ResumeGen(self.config)
        input_text = "Senior Software Engineer, Example"
        result = generator.remove_hyperlink(input_text)
        self.assertEqual(result, input_text)

    def test_process_text_with_hyperlinks(self):
        """Test processing text with hyperlinks."""
        generator = ResumeGen(self.config)
        mock_paragraph = Mock(spec=Paragraph)
        mock_paragraph.part = MagicMock()
        mock_paragraph.part.relate_to.return_value = "rId1"  # Return a string for r_id
        mock_paragraph.add_run.return_value = MagicMock()

        text = "Test with <a href='https://test.com'>Test Link</a>."
        with patch("docx.oxml.shared.OxmlElement") as mock_oxml:
            generator.process_text_with_hyperlinks(mock_paragraph, text)
            mock_paragraph.add_run.assert_called()
            mock_paragraph.part.relate_to.assert_called_once_with(
                "https://test.com", docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
            )
            mock_oxml.assert_any_call("w:hyperlink")  # Check for w:hyperlink among calls

    def test_generate_resume_structure(self):
        """Test generating resume document structure."""
        generator = ResumeGen(self.config)
        # Mock Document methods without patching read-only attributes
        generator.doc = Mock()
        generator.doc.add_heading = Mock()
        generator.doc.add_paragraph = Mock(return_value=Mock(spec=Paragraph))
        generator.doc.add_paragraph.return_value.part = MagicMock()
        generator.doc.add_paragraph.return_value.part.relate_to.return_value = "rId1"  # Return a string
        generator.doc.save = Mock()

        with patch("resumegen.logger") as mock_logger:
            generator.generate_resume(self.resume_data, "output.docx")

            # Verify headings
            generator.doc.add_heading.assert_any_call("Nicholas Borbaki", 0)
            generator.doc.add_heading.assert_any_call("Senior Software Engineer", 3)
            generator.doc.add_heading.assert_any_call("Experience", 1)
            generator.doc.add_heading.assert_any_call("Senior Software Engineer, Example - (2021-Present)", 2)

            # Verify paragraph for bullet point
            generator.doc.add_paragraph.assert_called_once_with(style="List Bullet")

            # Verify save
            generator.doc.save.assert_called_once_with("output.docx")
            mock_logger.info.assert_called_once_with("Resume saved to output.docx")

    def test_generate_resume_invalid_data(self):
        """Test generating resume with invalid data."""
        generator = ResumeGen(self.config)
        invalid_data = {"header": {}, "contents": []}  # Missing name and title
        with self.assertLogs("resumegen", level="ERROR") as log:
            with self.assertRaises(KeyError):
                generator.generate_resume(invalid_data, "output.docx")
            self.assertIn("Failed to generate resume", log.output[0])

if __name__ == "__main__":
    unittest.main()
