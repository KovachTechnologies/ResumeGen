#!/usr/bin/env python3

"""
ResumeGen: A tool to convert JSON resumes into formatted Word documents.
"""

import argparse
import datetime
import json
import logging
import re
from pathlib import Path
from typing import Dict, Any, Optional

import docx
import requests
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)

class ResumeGen:
    """Class to generate a Word document resume from JSON data."""
    
    def __init__(self, config: Dict[str, Any]) -> None:
        """Initialize ResumeGen with configuration settings."""
        self.config = config
        self.doc = docx.Document()
        self._apply_document_styles()

    def _apply_document_styles(self) -> None:
        """Apply default styles to the document."""
        for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
            style = self.doc.styles[style_name]
            font = style.font
            font.name = self.config["font"]["name"]
            font.size = Pt(self.config["font"]["size"])

        # Configure page margins
        for section in self.doc.sections:
            section.top_margin = Cm(self.config["margins"]["top"])
            section.bottom_margin = Cm(self.config["margins"]["bottom"])
            section.left_margin = Cm(self.config["margins"]["left"])
            section.right_margin = Cm(self.config["margins"]["right"])

    def add_hyperlink(self, paragraph: docx.text.paragraph.Paragraph, text: str, url: str) -> None:
        """Add a hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        new_run = docx.oxml.shared.OxmlElement("w:r")
        r_pr = docx.oxml.shared.OxmlElement("w:rPr")
        new_run.append(r_pr)
        new_run.text = text
        hyperlink.append(new_run)

        run = paragraph.add_run()
        run._r.append(hyperlink)
        run.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        run.font.underline = True

    def process_text_with_hyperlinks(self, paragraph: docx.text.paragraph.Paragraph, text: str) -> None:
        """Process text, converting HTML <a> tags to Word hyperlinks."""
        pattern = r'<a\s+[^>]*href=[\'"](.*?)[\'"][^>]*>(.*?)</a>'
        last_pos = 0
        for match in re.finditer(pattern, text):
            start, end = match.span()
            # Add text before the hyperlink
            if start > last_pos:
                paragraph.add_run(text[last_pos:start])
            # Add the hyperlink
            self.add_hyperlink(paragraph, match.group(2), match.group(1))
            last_pos = end
        # Add remaining text
        if last_pos < len(text):
            paragraph.add_run(text[last_pos:])

    def remove_hyperlink(self, text: str) -> str:
        """Remove HTML <a> tags from text, keeping the inner content."""
        return re.sub(r'<a\s+[^>]*>(.*?)</a>', r'\1', text)

    def generate_resume(self, resume_data: Dict[str, Any], output_file: str) -> None:
        """Generate a resume from JSON data and save it as a Word document."""
        try:
            # Add header
            self.doc.add_heading(resume_data["header"]["name"], 0)
            self.doc.add_heading(resume_data["header"]["title"], 3)

            # Sort and process content sections
            contents = sorted(resume_data["contents"], key=lambda x: x["id"], reverse=True)
            for datum in contents:
                self.doc.add_heading(datum["title"], 1)
                items = sorted(datum["content"], key=lambda x: x["id"], reverse=True)
                for item in items:
                    if item["position"]:
                        # Strip hyperlinks from position before rendering as heading
                        heading = self.remove_hyperlink(item["position"])
                        if item["date"]:
                            heading = f"{heading} - ({item['date']})"
                        self.doc.add_heading(heading, 2)

                    for bullet_item in item["items"]:
                        paragraph = self.doc.add_paragraph(style="List Bullet")
                        self.process_text_with_hyperlinks(paragraph, bullet_item)

            # Save the document
            self.doc.save(output_file)
            logger.info(f"Resume saved to {output_file}")
        except Exception as e:
            logger.error(f"Failed to generate resume: {str(e)}")
            raise

def fetch_json(url: str, headers: Dict[str, str]) -> Dict[str, Any]:
    """Fetch JSON data from a URL."""
    try:
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        return response.json()
    except requests.RequestException as e:
        logger.error(f"Failed to fetch JSON from {url}: {str(e)}")
        raise

def load_json(file_path: Path) -> Dict[str, Any]:
    """Load JSON data from a file."""
    try:
        with file_path.open("r") as f:
            return json.load(f)
    except (IOError, json.JSONDecodeError) as e:
        logger.error(f"Failed to load JSON from {file_path}: {str(e)}")
        raise

def main() -> None:
    """Main function to run ResumeGen."""
    parser = argparse.ArgumentParser(
        description="ResumeGen: Convert JSON resumes to Word documents.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    parser.add_argument(
        "--url",
        type=str,
        default=None,
        help="URL to fetch JSON resume data",
    )
    parser.add_argument(
        "--file",
        type=Path,
        default=None,
        help="Path to local JSON resume file",
    )
    parser.add_argument(
        "--output",
        type=str,
        default="{datetime}_resume.docx",
        help="Output Word document path (supports {datetime} placeholder)",
    )
    args = parser.parse_args()

    # Validate arguments
    if args.file is None and args.url is None:
        parser.error("Either --file or --url must be provided")
    if args.file is not None and args.url is not None:
        parser.error("Only one of --file or --url can be provided")

    # Configuration settings
    config = {
        "font": {"name": "Arial", "size": 8},
        "margins": {"top": 0.5, "bottom": 0.5, "left": 1.0, "right": 1.0},
    }

    # Replace datetime placeholder in output path
    output_file = args.output.format(datetime=datetime.datetime.now().strftime("%Y-%m"))

    try:
        # Load resume data
        headers = {
            "User-Agent": (
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_10_1) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/39.0.2171.95 Safari/537.36"
            )
        }
        if args.file:
            resume_data = load_json(args.file)
        else:
            resume_data = fetch_json(args.url, headers)

        # Generate resume
        generator = ResumeGen(config)
        generator.generate_resume(resume_data, output_file)
    except Exception as e:
        logger.critical(f"Application failed: {str(e)}")
        exit(1)

if __name__ == "__main__":
    main()
